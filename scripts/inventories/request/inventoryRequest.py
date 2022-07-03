"""
 # administrare - web platform for internal data management
 # Copyright (C) 2022 astrantialabs
 #
 # This program is free software: you can redistribute it and/or modify
 # it under the terms of the GNU General Public License as published by
 # the Free Software Foundation, either version 3 of the License, or
 # (at your option) any later version.
 #
 # This program is distributed in the hope that it will be useful,
 # but WITHOUT ANY WARRANTY; without even the implied warranty of
 # MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
 # GNU General Public License for more details.
 #
 # You should have received a copy of the GNU General Public License
 # along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""

"""
 # @fileoverview The InventoryRequest file.
 # @author Rizky Irswanda <rizky.irswanda115@gmail.com>
"""

import os
import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dependency import Dependency
from excel import Excel
from database import Database
from utility import Utility


class InventoryRequest():
    def writeRaw(currentDate):
        filePath = f"../{Dependency.inventoryRequestFolderPath}/Mentah {currentDate}.xlsx"

        Excel.create_file(filePath)
        workbook = Excel(filePath)
        workbook.change_sheet_name("Sheet", "Seluruh")
        workbook.set_zoom(85)

        requestCollection = Database.getCollection(Dependency.mongoDBURI, Dependency.databaseInventory, Dependency.collectionInventoryRequest)
        inventoryRequestDocument = requestCollection.find_one({"tahun": 2022})

        masterCollection = Database.getCollection(Dependency.mongoDBURI, Dependency.databaseInventory, Dependency.collectionInventoryMaster)
        inventoryMasterDocument = masterCollection.find_one({"tahun": 2022})

        userData = Utility.readJSON("./json/request_user_data.json")

        InventoryRequest.writeHeaderRaw(workbook)
        InventoryRequest.writeMainRaw(workbook, inventoryRequestDocument, inventoryMasterDocument)

        for userObject in userData:
            workbook.create_sheet(userObject.get("username"))
            workbook.change_sheet(userObject.get("username"))
            workbook.set_zoom(85)

            requestData = { "barang": [] }

            for dateObject in userObject.get("date"):
                for requestObject in dateObject.get("request"):
                    requestData.get("barang").append(requestObject)


            InventoryRequest.writeHeaderRaw(workbook)
            InventoryRequest.writeMainRaw(workbook, requestData, inventoryMasterDocument)

            for dateObject in userObject.get("date"):
                dateName = f'{userObject.get("username")} {"".join(dateObject.get("date").split("-"))}'
                workbook.create_sheet(dateName)
                workbook.change_sheet(dateName)
                workbook.set_zoom(85)

                InventoryRequest.writeHeaderRaw(workbook)
                InventoryRequest.writeMainRaw(workbook, { "barang": dateObject.get("request") }, inventoryMasterDocument)
                    

        workbook.save()


    def writeHeaderRaw(workbook):
        workbook.write_value_multiple("A1", "J1", ["No.", "Peminta", "Kategori", "Barang", "Jumlah", "Satuan", "Keterangan", "Dibuat", "Direspon", "Status"])
        workbook.font_multiple("A1", "J1", size = 12, bold = True)
        workbook.alignment_multiple("A1", "J1", vertical = "center", horizontal = "center")
        workbook.border_multiple("A1", "J1", "all", style = "thin")

    
    def writeMainRaw(workbook, requestData, masterData):
        rowCount = 2
        for requestItemIndex, requestItemObject in enumerate(requestData.get("barang")):
            for masterCategoryObject in masterData.get("kategori"):
                if(requestItemObject.get("kategori_id") == masterCategoryObject.get("id")):
                    for masterItemObject in masterCategoryObject.get("barang"):
                        if(requestItemObject.get("barang_id") == masterItemObject.get("id")):
                            status = Utility.convertStatus(requestItemObject.get("status"))

                            mainValue = [
                                requestItemIndex + 1,
                                requestItemObject.get("username"),
                                masterCategoryObject.get("kategori"),
                                masterItemObject.get("nama"),
                                requestItemObject.get("total"),
                                masterItemObject.get("satuan"),
                                requestItemObject.get("deskripsi"),
                                requestItemObject.get("created_at"),
                                requestItemObject.get("responded_at"),
                                status
                            ]
            
                            workbook.write_value_multiple(["A", rowCount], ["J", rowCount], mainValue)
                            workbook.border_multiple(["A", rowCount], ["J", rowCount], "all", style="thin")
                            workbook.alignment_multiple(["A", rowCount], ["J", rowCount], vertical="center", horizontal="center")

                            workbook.alignment_multiple(["B", rowCount], ["D", rowCount], vertical="center", horizontal="left")
                            workbook.alignment_multiple(["G", rowCount], ["I", rowCount], vertical="center", horizontal="left")

                            rowCount += 1


        workbook.adjust_width("A1", ["J", rowCount], extra_width=2)


    def writeUser(userId, dateId):
        masterCollection = Database.getCollection(Dependency.mongoDBURI, Dependency.databaseInventory, Dependency.collectionInventoryMaster)
        inventoryMasterDocument = masterCollection.find_one({"tahun": 2022})

        optionData = Utility.readJSON("./json/request_option_data.json")
        userData = Utility.readJSON("./json/request_user_data.json")

        for userObject in optionData:
            if(userObject.get("id") == userId):
                usernameValue = userObject.get("name")

                for dateObject in userObject.get("date"):
                    if(dateObject.get("id") == dateId):
                        dateValue = dateObject.get("date")


        splittedDate = Utility.slugifyDate(dateValue).split("-")

        fileTemplate = "./media/User Request Template.docx"
        filePath = f"../{Dependency.inventoryRequestFolderPath}/{usernameValue} {Utility.slugifyDate(dateValue)}.docx"

        document = Document(fileTemplate)

        for paragraph in document.paragraphs:
            if("Balikpapan" in paragraph.text):
                paragraph.add_run(f" {splittedDate[2]} {Utility.translateMonthName(Utility.convertNumberToMonthName(splittedDate[1]))} {splittedDate[1]}")

        
        document.paragraphs[len(document.paragraphs) - 2].add_run(f"{' ' * 90} {usernameValue}")

        for userObject in userData:
            if(userObject.get("username") == usernameValue):
                for dateObject in userObject.get("date"):
                    if(dateObject.get("date") == dateValue.split(" ")[0]):

                        requestCount = 0
                        for requestIndex, requestObject in enumerate(dateObject.get("request")):
                            requestCount += 1


                        while requestCount > len(document.tables[0].rows) - 1:
                            document.tables[0].add_row()

                        rowCount = 1
                        for requestIndex, requestObject in enumerate(dateObject.get("request")):
                            for categoryObject in inventoryMasterDocument.get("kategori"):
                                if(categoryObject.get("id") == requestObject.get("kategori_id")):
                                    for itemObject in categoryObject.get("barang"):
                                        if(itemObject.get("id") == requestObject.get("barang_id")):
                                            itemName = itemObject.get("nama")
                                            itemUnit = itemObject.get("satuan")


                            row = document.tables[0].rows[rowCount]

                            value = [
                                requestIndex + 1,
                                itemName,
                                requestObject.get("total"),
                                itemUnit,
                                requestObject.get("deskripsi")
                            ]

                            for i in range(len(value)):
                                row.cells[i].text = str(value[i])
      
                                for paragraph in row.cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)
                            
                            
                            cellCenterList = [0, 2, 3]
                            for cell in cellCenterList:
                                for paragraph in row.cells[cell].paragraphs:
                                    paragraph.alignment= WD_ALIGN_PARAGRAPH.CENTER


                            rowCount += 1


        document.save(filePath)


    def updateUserData():
        requestCollection = Database.getCollection(Dependency.mongoDBURI, Dependency.databaseInventory, Dependency.collectionInventoryRequest)
        inventoryRequestDocument = requestCollection.find_one({"tahun": 2022})

        usernameArray = []
        for requestMain in inventoryRequestDocument.get("barang"):
            if(requestMain.get("status") == 1):
                if(requestMain.get("username") not in usernameArray):
                    usernameArray.append(requestMain.get("username"))


        userData = []
        for usernameIndex, usernameItem in enumerate(usernameArray):
            newUserObject = {
                "id": usernameIndex + 1,
                "username": usernameItem,
                "date": []
            }

            userData.append(newUserObject)


        for userObject in userData:
            dateArray = []
            for requestMain in inventoryRequestDocument.get("barang"):
                if(requestMain.get("status") == 1):
                    if(requestMain.get("username") == userObject.get("username")):
                        createdAt = (requestMain.get("created_at").split(" "))[0]
                        
                        if(createdAt not in dateArray):
                            dateArray.append(createdAt)


            for dateIndex, dateItem in enumerate(dateArray):
                newDateObject = {
                    "id": dateIndex + 1,
                    "date": dateItem,
                    "request": []
                }

                userObject.get("date").append(newDateObject)

        
        for requestMain in inventoryRequestDocument.get("barang"):
            if(requestMain.get("status") == 1):
                for userObject in userData:
                    if(requestMain.get("username") == userObject.get("username")):
                        for dateObject in userObject.get("date"):
                            if((requestMain.get("created_at").split(" "))[0]) == dateObject.get("date"):
                                newRequestObject = requestMain
                                newRequestObject["internal_id"] = len(dateObject.get("request")) + 1

                                dateObject.get("request").append(newRequestObject)


        Utility.writeJSON("./json/request_user_data.json", userData) 


    def updateOptionData():
        files = os.listdir("../spreadsheets/inventories/request")

        files.remove(".gitkeep")
        
        fileOptionArray = [["Mentah", [["Terbaru", True]]]]

        jsonUserRequstData = Utility.readJSON("./json/request_user_data.json")
        for userObject in jsonUserRequstData:
            
            dateArray = []
            for dateObject in userObject.get("date"):
                dateArray.append([f"{dateObject.get('date')} 00:00:00", True])

            
            fileOptionArray.append([userObject.get("username"), dateArray])
        

        for file in files:
            fileNameArray = file.split(".")[0].split(" ")
            if(len(fileNameArray) > 1):
                fileDate = fileNameArray[-1].split("-")
                formattedFileDate = f"{fileDate[0]}-{fileDate[1]}-{fileDate[2]} {fileDate[3]}:{fileDate[4]}:{fileDate[5]}"

                dateIsValid = None
                try:
                    dateValidation = datetime.datetime(int(fileDate[0]), int(fileDate[1]), int(fileDate[2]), int(fileDate[3]), int(fileDate[4]), int(fileDate[5]))
                    dateIsValid = True

                except ValueError:
                    dateIsValid = False

                if(dateIsValid):
                    fileName = " ".join(fileNameArray[0:-1])

                    if(fileName in ["Mentah"]):
                        fileNameIsValid = True
                        for fileItem in fileOptionArray:
                            if(fileItem[0] == fileName):
                                fileNameIsValid = False

                        
                        if(fileNameIsValid):
                            fileOptionArray.append([fileName, [[formattedFileDate, False]]])

                        elif(not fileNameIsValid):
                            for fileItem in fileOptionArray:
                                if(fileItem[0] == fileName):

                                    fileItemArrayIsValid = True
                                    for fileItemArray in fileItem[1]:
                                        if(fileItemArray[0] == formattedFileDate):
                                            fileItemArrayIsValid = False
                                    
                                    if(fileItemArrayIsValid):
                                        fileItem[1].append([formattedFileDate, False])


        fileOptionData = []
        for fileItemIndex, fileItem in enumerate(fileOptionArray):
            fileDateData = []
            for fileDateIndex, fileDate in enumerate(fileItem[1]):
                newFileDateObject = {
                    "id": fileDateIndex + 1,
                    "date": fileDate[0],
                    "creatable": fileDate[1]
                }

                fileDateData.append(newFileDateObject)

            
            newFileOptionObject = {
                "id": fileItemIndex + 1,
                "name": fileItem[0],
                "date": fileDateData
            }

            fileOptionData.append(newFileOptionObject)

        
        Utility.writeJSON("./json/request_option_data.json", fileOptionData)
